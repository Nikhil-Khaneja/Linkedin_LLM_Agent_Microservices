from __future__ import annotations

from io import BytesIO
from typing import Any
import re

# Used by extract_keywords (and shared with ai matching); keep in sync with orchestrator heuristics.
_STOPWORDS = frozenset({
    'the', 'and', 'for', 'with', 'from', 'that', 'this', 'will', 'role', 'team', 'your', 'you', 'our', 'are', 'job', 'work',
    'about', 'have', 'has', 'into', 'who', 'years', 'year', 'plus', 'using', 'use', 'required', 'preferred', 'strong', 'ability',
    'experience', 'skills', 'candidate', 'position', 'including', 'across', 'such', 'their', 'they', 'them', 'open', 'closed',
})


def extract_keywords(text: str, limit: int = 12) -> list[str]:
    """Token-level keyword skim for resumes/jobs (no LLM). Safe for member_profile and shared code."""
    tokens = re.findall(r'[A-Za-z][A-Za-z0-9+#./-]{2,}', (text or '').lower())
    seen: set[str] = set()
    out: list[str] = []
    for token in tokens:
        if token in _STOPWORDS:
            continue
        if token not in seen:
            seen.add(token)
            out.append(token)
        if len(out) >= limit:
            break
    return out


def _normalize_text(text: str) -> str:
    text = text.replace('\x00', ' ')
    text = re.sub(r'\r\n?', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()


def extract_text_from_bytes(filename: str | None, content: bytes | None, content_type: str | None = None) -> str:
    if not content:
        return ''
    filename = (filename or '').lower()
    content_type = (content_type or '').lower()

    if filename.endswith('.docx') or 'wordprocessingml' in content_type or 'officedocument' in content_type:
        try:
            from docx import Document

            doc = Document(BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = (cell.text or '').strip()
                        if t:
                            parts.append(t)
            return _normalize_text('\n'.join(parts))
        except Exception:
            pass

    if filename.endswith('.pdf') or 'pdf' in content_type:
        try:
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(content))
            text_parts = []
            for page in reader.pages:
                try:
                    text_parts.append(page.extract_text() or '')
                except Exception:
                    continue
            return _normalize_text('\n'.join(text_parts))
        except Exception:
            return ''

    for encoding in ('utf-8', 'latin-1'):
        try:
            return _normalize_text(content.decode(encoding, errors='ignore'))
        except Exception:
            continue
    return ''


def extract_text_from_response(url: str, response_content: bytes, content_type: str | None = None) -> str:
    return extract_text_from_bytes(url.rsplit('/', 1)[-1], response_content, content_type)
